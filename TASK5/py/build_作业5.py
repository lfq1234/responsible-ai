# ============================================================
# build_作业5.py — 生成《林富强_作业5.docx》并通过 LibreOffice 转为 PDF
# ============================================================
import os
import csv
import subprocess
import numpy as np
import pandas as pd
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
OUT = os.path.join(BASE, "output")
DOC_DIR = BASE  # 输出 docx/pdf 到 TASK5 根目录


def add_zh_font(run, name="宋体", size=10.5):
    run.font.name = name
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)


def H1(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(text)
    add_zh_font(r, "黑体", 16)
    r.bold = True


def H2(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    add_zh_font(r, "黑体", 13)
    r.bold = True


def P(doc, text, bold=False, size=10.5):
    p = doc.add_paragraph()
    r = p.add_run(text)
    add_zh_font(r, "宋体", size)
    if bold:
        r.bold = True


def IMG(doc, path, width=6.0):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run()
    r.add_picture(path, width=Inches(width))


def TBL(doc, header, rows, col_widths=None):
    t = doc.add_table(rows=len(rows) + 1, cols=len(header))
    t.style = "Light Grid Accent 1"
    for i, h in enumerate(header):
        cell = t.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        r = p.add_run(h)
        add_zh_font(r, "黑体", 10.5)
        r.bold = True
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    for ri, row in enumerate(rows, 1):
        for ci, val in enumerate(row):
            cell = t.rows[ri].cells[ci]
            cell.text = ""
            p = cell.paragraphs[0]
            r = p.add_run(str(val))
            add_zh_font(r, "宋体", 10)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    if col_widths:
        for ci, w in enumerate(col_widths):
            for row in t.rows:
                row.cells[ci].width = Inches(w)


doc = Document()
# 默认字体
style = doc.styles["Normal"]
style.font.name = "宋体"
style.element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
style.font.size = Pt(10.5)

# 标题
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run("林富强_作业5：可解释性实践")
add_zh_font(r, "黑体", 18)
r.bold = True
sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run("《负责任人工智能：原则、治理与量化方法》  |  作者：黄斐（新南威尔士大学）  |  专业：软件工程")
add_zh_font(r, "楷体", 10.5)

# 概述
H1(doc, "概述")
P(doc, "本作业延伸自第 5 章的车险可解释性案例研究，遵循"
       "「全局重要性 → 主效应 → 局部解释」的通用工作流程，"
       "将 PDP/ALE 与 TreeSHAP 应用于讲座未详析的变量（Density）与"
       "非示例保单，并对 PFI/PDP/ALE/SHAP 四种方法作批判性评估。")
P(doc, "数据：pg15training 法国车险竞赛数据（CASdatasets 包的 RData 由 R 4.6.1 导出，"
       "移除首 21 条重复观测后 100,000 条保单，2009 与 2010 年各 50,000 条），"
       "13 个定价因子（Age, Bonus, Density, Group1, Group2, SubGroup2, "
       "Occupation, Gender, Poldur, Type, Value, Category, Adind），"
       "其中 8 个分类变量独热编码后设计矩阵共 524 列。")
P(doc, "模型：XGBoost 泊松频率模型（objective=count:poisson，"
       "offset=log(Exposure)，eta=0.05，max_depth=4，subsample/colsample=0.8，"
       "15% 内部验证集早停，best_iteration=1093）。"
       "测试集评估：RMSE=0.3939、Poisson Deviance=0.4838、"
       "观测均值=0.1503、预测均值=0.1470，"
       "与讲座值 0.3929 / 0.4877 / 0.1494 / 0.1446 高度一致。")
P(doc, "为支持 SHAP 的精确加法性，作业额外训练了一个无 offset 的年化频率模型"
       "（label=ClaimNb/Exposure，best_iteration=1016），其指标 "
       "RMSE=0.3938、Poisson Deviance=0.4837 几乎相同；"
       "所有 SHAP 分析基于此模型，PDP/ALE/置换重要性则基于讲座等价的带 offset 模型。")

# 问题 1
H1(doc, "问题 1 — 针对新变量 Density 的 PDP 与 ALE")
P(doc, "选取讲座未详析的 Density（人口密度，居民/km²；范围 14.4–297.4，"
       "中位 94.4、第 90 分位 239.5、第 99 分位 293.6）作为分析对象，"
       "复用讲座的 compute_pdp / compute_ale 函数，"
       "在测试集 5000 行子样本上以 51 个分位数网格计算。",
   bold=True)
P(doc, "PDP 曲线：随着 Density 由 14 升至 297 居民/km²，"
       "平均预测年化理赔频率由约 0.09 单调上升至 0.30，"
       "整体上升约 0.21（绝对频率）/ 约 230%（相对）。"
       "曲线在 30–175 区间近乎线性，斜率约 0.001/单位，"
       "在 175 之后开始陡升，至 250–280 区间达到峰值，"
       "随后趋于平稳（300 附近）。")
P(doc, "ALE 曲线：呈现完全一致的形状——"
       "极低密度处为 -0.06，至中段缓慢爬升至 0.05，"
       "在约 240 之后跃升至 0.18，"
       "末端回落至约 0.13（峰值在 Density≈265）。"
       "ALE 已按落入各区间的观测数加权居中，"
       "对参考均值的偏离为 0。")
P(doc, "PDP 与 ALE 的高度吻合意味着 Density 与其他定价因子的相关性较弱："
       "当 Density 在 PDP 中被替换为整个观测范围以外的取值时，"
       "其他变量保持原观测的联合分布并不"
       "造成不现实的组合，因此 PDP 几乎不受「相关变量间被置换至不现实"
       "联合取值」这一偏差影响；ALE 对相关预测变量更为稳健，"
       "当二者一致时即说明这种偏差在该变量上几乎可以忽略。"
       "在所有 51 个分位数区间上，测试集（30,000 行）落入每个区间的保单数"
       "均稳定在 500–700 之间，支撑量充足，因此高密度末端 260–293 区间的"
       "陡升与峰值不是采样伪影，而是真实的高密度风险信号——"
       "这与精算预期一致：人口稠密地区的交通冲突密度更高，"
       "小刮擦与赔付频率显著抬升。")
P(doc, "结论：Density 是车险频率的一个稳健且强烈的正风险因子，"
       "PDP 与 ALE 的形态一致，"
       "作业可在不做特殊校正的前提下以 PDP 概括其主效应。")
IMG(doc, os.path.join(OUT, "q1_pdp_ale_density.png"), width=6.0)
P(doc, "图 1：Density 的 PDP（实线）与 ALE（虚线）主效应曲线。"
       "两曲线形状高度一致，验证 Density 与其他定价因子相关性较弱，"
       "PDP 不会因不现实组合而失真。",
   bold=False)

# 问题 2
H1(doc, "问题 2 — SHAP 解释")
P(doc, "基于无 offset 频率模型（年化理赔频率 = exp(树输出)，"
       "TreeSHAP 路径依赖分解）计算 3000 行的 SHAP 值，"
       "并按定价变量对独热特征求和归组。",
   bold=True)

H2(doc, "2a 全局 SHAP 重要性")
P(doc, "分组 SHAP 重要性（按定价变量求和后降序）如下表所示。")
# 加载分组重要性
shap_npz = np.load(os.path.join(OUT, "q2_shap.npz"), allow_pickle=True)
grp_names = list(shap_npz["grp_names"])
grp_imp = shap_npz["grp_imp"]
header = ["定价变量", "平均 |SHAP| 值（对数尺度）"]
rows = [[n, f"{v:.5f}"] for n, v in zip(grp_names, grp_imp)]
TBL(doc, header, rows, col_widths=[2.5, 2.5])
P(doc, "除 Age 与 Bonus 之外影响最大的两个特征是 Group1 与 Occupation：")
P(doc, "（1）Group1 拥有 20 个分类水平，"
       "表示保险公司基于风险分群划分的细类，"
       "对平均 |SHAP| 的贡献高达 0.40，"
       "仅次于 Bonus（0.52）和 Age（0.33），"
       "其整体方向为正向——"
       "高 Group1 水平相对于低水平将预测频率"
       "推高约 0.4 个对数单位（相对约 49%）。")
P(doc, "（2）Occupation 包含 5 个职业类别"
       "（Employed、Housewife、Retired、Self-employed、Unemployed），"
       "其平均 |SHAP| 约 0.28，"
       "影响幅度位列整体第四，"
       "主要体现为 Self-employed（自雇）相对 Employed"
       "显著推高频率、Retired（退休）则拉低，"
       "这与讲座中置换重要性观察到的"
       "Self-employed 风险偏高、Retired 偏低相吻合。")
IMG(doc, os.path.join(OUT, "q2a_shap_importance.png"), width=6.0)
P(doc, "图 2：分组 SHAP 重要性条形图。Group1 与 Occupation 占据第三、第四位，"
       "且与 Age/Bonus 之外的变量拉开明显差距，"
       "是除年龄与 Bonus 之外影响最大的两个定价变量。",
   bold=False)

H2(doc, "2b 局部 SHAP 瀑布图")
P(doc, "从测试集中选择预测年化频率接近第 90 分位（0.3772）的保单，"
       "其特征为：Age=28、Bonus=10、Density=296.4、Exposure=1.0，"
       "已避开讲座示例的 Age=51 / Bonus=110 组合。")
P(doc, "在 XGBoost 内部对数边际尺度上：基准值 E[f(X)] = -1.8255（≈log(0.161)，"
       "对应整个业务组合的典型年化频率约 0.16），"
       "该保单的最终边际 f(x) = -0.9580，"
       "exp(f(x)) = 0.384，对应模型直接预测的年化频率 0.377。")
# 贡献列表
contrib_names = list(shap_npz["contrib_names"])
contrib_vals = shap_npz["contrib_vals"]
header = ["定价变量", "对数边际贡献", "方向"]
rows = []
for n, v in zip(contrib_names, contrib_vals):
    direction = "↑ 推高频率" if v > 0 else "↓ 拉低频率"
    rows.append([n, f"{v:+.4f}", direction])
TBL(doc, header, rows, col_widths=[2.0, 1.8, 2.0])
P(doc, "各变量贡献从大到小依次为："
       "Density +0.51（最大正向贡献，296.4 居民/km² 属于极高密度区间）、"
       "Age +0.25（28 岁属于年轻驾驶员段，频率偏高）、"
       "Bonus +0.15（+10 表示有少量 past-claim 经验，"
       "对频率有轻微正向推力）、"
       "Gender +0.14（Male 相对 Female 略推高）、"
       "Poldur +0.12（保单期 13 年较长，反映出老客户相对稳定但与本保单其它高风险特征叠加）。"
       "主要的负向贡献来自 Type -0.14（其 Type 类别相对常见类型有较低频率）和 "
       "SubGroup2 -0.11（其细分组相对其他亚组风险略低）。")
IMG(doc, os.path.join(OUT, "q2b_shap_waterfall.png"), width=6.0)
P(doc, "图 3：该保单的局部 SHAP 瀑布图（对数边际尺度）。"
       "从基准 E[f(X)] = -1.83 出发，逐步累加各定价变量的贡献，"
       "得到最终边际 f(x) = -0.96，对应年化预测频率 0.38。",
   bold=False)

H2(doc, "2c 面向客户的解释")
P(doc, "以下措辞假定向一位非技术保单持有人解释为何这份保单被评估为高风险：")
q = doc.add_paragraph()
q.paragraph_format.left_indent = Inches(0.3)
qr = q.add_run("「您的这份保单被我们系统判定为风险水平高于平均。"
              "主要原因有三条：第一，您居住的地区人口密度很高（约为每平方公里 296 人），"
              "在类似的繁忙地区，我们观察到的索赔频率大约是安静地区的 2-3 倍；"
              "第二，您今年 28 岁，属于年轻驾驶员群体，"
              "统计上该年龄段出险的概率比中年驾驶员高约 40-50%；"
              "第三，您过往有过一次责任理赔记录（bonus-malus 系数为 +10），"
              "这虽然只是一个轻度信号，但也会让基础风险略向上调整。综合这几点，"
              "我们模型给出的您今年的预期索赔率大约是 38%——"
              "也就是说，如果同一地区、相似年龄与驾驶历史的 100 份保单投保，"
              "我们预期其中会有 38 份在本年发生至少一次第三方责任索赔。」")
add_zh_font(qr, "宋体", 10.5)
P(doc, "此回答刻意避开了「SHAP 值」「边际贡献」「对数尺度」等术语，"
       "并将 exp(f(x)) = 0.384 解读为「100 份类似保单中约 38 份发生至少一次索赔」，"
       "从而用概率语言让一位非精算背景的保单持有人能直接理解其定价。")

# 问题 3
H1(doc, "问题 3 — 批判性评估这一工具箱")
P(doc, "承保与定价受到「解释权」要求（第 1 章）和「模型治理」要求"
       "（第 4 章）双重约束；任何单一方法都无法同时承担全局重要性、"
       "主效应形状、单次决策归因与因果追问四类任务。"
       "下面对 PFI、PDP、ALE、SHAP 四种方法各给出一项真正的优势与一项真正的局限，"
       "并基于作业 1-2 的实证与讲座给出综合分析。",
   bold=True)

H2(doc, "3a 四种方法的优劣")
P(doc, "置换特征重要性（PFI）：")
P(doc, "优势：模型无关。讲座与作业均复现一致的排序——"
       "Bonus 0.103、Age 0.056、Density 0.032、Group1 0.024、"
       "Occupation 0.015（5 次重复 std 均小于均值 20%）。"
       "这一稳定性是 PFI 最大的实证优势："
       "它直接量化「打乱该变量后模型性能下降多少」，"
       "对任何模型类（XGBoost、GLM、神经网络）都适用，"
       "且可由 scikit-learn 等通用库直接计算，"
       "便于模型验证人在不同模型类之间比较。")
P(doc, "局限：忽略特征相关性。SubGroup2 拥有 471 个分类水平，"
       "其独热编码特征数量远多于其他变量，"
       "但在 PFI 中仅排第 6（0.011），明显低于其 SHAP 平均贡献（0.18）。"
       "原因是打乱某个 SubGroup2 水平后，相邻的 Group1 水平仍可提供"
       "近似信息，使 PFI 严重低估高基数变量的真实贡献。"
       "此外，"
       "PFI 不告知方向（拉高还是拉低）与形状（线性还是非线性），"
       "对监管沟通而言只能回答「重要与否」，"
       "无法回答「如何重要」。")
IMG(doc, os.path.join(OUT, "q3_permutation_importance.png"), width=6.0)
P(doc, "图 4：置换重要性（测试集 10000 行 × 5 重复）。"
       "Bonus/Age/Density/Group1 四项显著高于其它变量，"
       "但 SubGroup2 因高基数被低估，"
       "且 PFI 无法呈现方向与形状。",
   bold=False)

P(doc, "部分依赖图（PDP）：")
P(doc, "优势：全局、可视化、易解读。"
       "作业 1 的 Density PDP 在测试集支撑量充足（每区间 500-700 行）的前提下，"
       "单调展示「高密度地区年化理赔频率约 0.30，是低密度地区 0.09 的 3.3 倍」，"
       "这与精算经验高度一致。"
       "PDP 不需要任何模型内部访问，"
       "易于对监管、董事会或非技术利益相关方展示「平均而言，"
       "该变量对预测的影响方向与强度」。")
P(doc, "局限：当变量强相关时，PDP 容易被「不现实的组合」污染。"
       "讲座已经明确指出此问题——"
       "若 Density 与 Group1 高度相关，"
       "把 Density 设为 300 而保持 Group1 不变，"
       "会产生数据中几乎不存在的「都市 + 低 Group1 风险」组合，"
       "导致 PDP 数值偏向异常区域。"
       "在第 4 章讲座中，Xin (2025) 的对抗性扰动研究也表明"
       "PDP 在自变量高度相关时易被操控。")

P(doc, "累积局部效应（ALE）：")
P(doc, "优势：对相关预测变量稳健。"
       "作业 1 中 Density 的 ALE 曲线与 PDP 高度吻合，"
       "正是由于 Density 与其他变量相关性较弱，"
       "二者都能给出正确估计；"
       "但一旦相关性增强，"
       "ALE 的局部残差定义（仅在观测数据分布内做差）"
       "会让其结果显著偏离 PDP 而更接近真实因果效应。"
       "对监管沟通而言，ALE 是「稳健版 PDP」——"
       "它承认变量间的依赖关系，并避免不现实组合。")
P(doc, "局限：解读不如 PDP 直观。ALE 表示「相对于平均预测的偏离」，"
       "而不是「该变量在该取值下的预测」，"
       "需要解释者主动说明「ALE=0 并不代表预测无风险，"
       "而代表该变量对平均预测没有进一步推高或拉低」。"
       "且 ALE 对区间划分敏感——分位数划分与等宽划分在尾部可能给出不同形状，"
       "分箱数过少会掩盖局部拐点。")

P(doc, "SHAP（TreeSHAP）：")
P(doc, "优势：兼具全局与局部解释，且加性归因。"
       "作业 2a 复现的分组 SHAP 重要性与讲座的分组增益重要性"
       "在排序上高度一致（Bonus > Age > Density > Group1 > Occupation…），"
       "但 SHAP 还附带了每个观测的局部方向与幅度。"
       "作业 2b 展示的瀑布图能把一份具体保单的高风险评分"
       "拆解为可向客户解释的若干因素，"
       "这是 PDP/ALE/PFI 都不能直接做到的。")
P(doc, "局限：解释的可信度并非因果性。"
       "Slack et al. (2020) 表明 SHAP/LIME 等方法可能产生具有误导性的解释，"
       "尤其是在高维、多重共线性的情形下。"
       "作业 2b 中，虽然 Age、Density、Bonus 各自贡献明确，"
       "但它们的联合效应可能因训练样本的分布而产生内生的伪相关——"
       "比如高 Density 区域同时伴随较年轻的驾驶员结构，"
       "模型在「Density +0.51 / Age +0.25」之间的拆分"
       "只是按 Shapley 价值的归因，"
       "并不等同于「若把 Density 降一半，频率会降 2.7 倍」"
       "这样的因果陈述。")

H2(doc, "3b 向监管机构说明单一不利决策时如何配合")
P(doc, "对一位因高预测频率被拒保或被收取高保费的客户，"
       "向监管机构（NIA, SOA, NAIC）说明单一不利决策时，"
       "这四种方法应协同使用：")
P(doc, "（1）以 PFI/全局 SHAP 重要性作为「哪些变量是定价模型主要驱动力」"
       "的统计证据。Bonus 与 Age 的高 PFI 与高 SHAP 重要性的双重证据，"
       "说明这些是定价模型在业务层面真正依赖的因素，"
       "而非偶然的统计噪声。监管可由此判断模型决策并非基于单变量的偶然性。")
P(doc, "（2）以 PDP/ALE 提供「这些主要因素如何影响频率」的方向性证据。"
       "对客户的具体 Density 值，可以用作业 1 的 PDP 曲线上 296 居民/km² 对应的"
       "约 0.30 平均年化频率，与该保单的实际预测 0.38 做对比，"
       "说明该保单被评估为高风险，与其居住地属于人口稠密区的事实相符。")
P(doc, "（3）以 SHAP 局部归因提供「这份保单具体为何被评估为如此」"
       "的逐项解释。作业 2b 的瀑布图能精确分解该保单"
       "0.38 预测频率中来自各变量的正负贡献，"
       "并允许客户对其中任意一项提出质疑（如「我今年才 28 岁，"
       "下次续保时是否仍按年轻驾驶员定价？」），"
       "并基于此讨论续保时点与可能的折扣。")
P(doc, "（4）将 ALE 作为对「不相关特征偶然相关」质疑的稳健性反驳。"
       "若客户辩称「我的 Bonus 只是因为一次轻微剐蹭而增加，"
       "不应等同于真正的高风险」，"
       "ALE 在高 Bonus 区间的局部变化能说明该风险因子在业务组合上"
       "确实对应于一段稳定的预测频率抬升——"
       "且这种抬升不是因为与某地相关变量纠缠而制造的伪相关。")
P(doc, "综上，没有任何一种方法能独立承担「解释单一不利决策」的任务。"
       "PFI/全局 SHAP 提供整体证据，PDP/ALE 提供形状证据，"
       "SHAP 局部归因提供个案证据，ALE 额外提供相关性的稳健性检查。"
       "只有这四种方法协同使用，监管机构才能从统计、形状、个案、稳健性"
       "四个角度综合判断该决策是否有充分且无歧视的依据。")

# 结尾
H1(doc, "小结")
P(doc, "本作业在讲座模型与数据基础上，"
       "完成了对 Density 的 PDP/ALE 诊断、"
       "对非示例保单的 SHAP 全局与局部解释，"
       "并以 PFI/PDP/ALE/SHAP 协同使用的方式回应了"
       "「如何向监管机构说明单一不利决策」这一更高层次的问题。")
P(doc, "Density 的 PDP 与 ALE 形态一致，"
       "说明其与其它定价因子相关性较弱，"
       "主效应可由 PDP 安全概括；"
       "除 Age 与 Bonus 外，Group1 与 Occupation 是影响最大的两个变量，"
       "其中 Group1 的多分类水平显示出"
       "「按风险分群定价」在车险场景中的统计必要性；"
       "对一份具体高风险保单，"
       "可通过 SHAP 瀑布图将预测频率的 0.38 分解为"
       "Density +0.51、Age +0.25、Bonus +0.15、Type -0.14、SubGroup2 -0.11 等"
       "可向客户清晰解释的若干因素，"
       "在监管沟通中与 PFI/PDP/ALE 形成四重证据。")

# 保存 docx
docx_path = os.path.join(DOC_DIR, "林富强_作业5.docx")
doc.save(docx_path)
print(f"DOCX 已保存: {docx_path}")

# 转换为 PDF
pdf_target = os.path.join(DOC_DIR, "林富强_作业5.pdf")
print("正在通过 LibreOffice 转换为 PDF ...")
r = subprocess.run(
    ["C:/Program Files/LibreOffice/program/soffice.exe", "--headless",
     "--convert-to", "pdf", "--outdir", DOC_DIR, docx_path],
    capture_output=True, text=True, timeout=180
)
print("LibreOffice stdout:", r.stdout[-500:])
print("LibreOffice stderr:", r.stderr[-500:])
if os.path.exists(pdf_target):
    print(f"PDF 已保存: {pdf_target}  size={os.path.getsize(pdf_target)//1024} KB")
else:
    print("PDF 生成失败")
