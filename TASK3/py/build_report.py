# ============================================================
# build_report.py — 生成《林富强_作业3.docx》报告
# 内容：讲义基线复现 + 六题（代码/结果/图表/书面回答）
# ============================================================

import os
import sys
import json
import numpy as np
import pandas as pd
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

PY = os.path.dirname(os.path.abspath(__file__))
OUT = os.path.join(PY, "output")
TASK3 = r"C:/Users/LENOVO/Desktop/responsible-ai/TASK3"

# ---------- 读取结果 ----------
fa = pd.read_csv(os.path.join(OUT, "fairness_accuracy.csv"))
mean_table = pd.read_csv(os.path.join(OUT, "mean_premium_table.csv"))
q1 = pd.read_csv(os.path.join(OUT, "q1_lambda_scan.csv"))
scan = pd.read_csv(os.path.join(OUT, "compas_theta_scan.csv"))
summary = json.load(open(os.path.join(OUT, "q1_q2_summary.json"), encoding="utf-8"))
compas = json.load(open(os.path.join(OUT, "compas_summary.json"), encoding="utf-8"))
q3 = json.load(open(os.path.join(OUT, "q3_dir.json"), encoding="utf-8"))

doc = Document()

# 全局中文字体
style = doc.styles["Normal"]
style.font.name = "Times New Roman"
style.font.size = Pt(11)
style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")


def add_heading(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = "Times New Roman"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
        run.font.color.rgb = RGBColor(0, 0, 0)
    return h


def add_para(text, bold=False, size=11, align=None):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    r.font.size = Pt(size)
    if align:
        p.alignment = align
    return p


def add_code(code):
    for line in code.strip("\n").split("\n"):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(line if line else " ")
        r.font.name = "Consolas"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        r.font.size = Pt(8.5)
        p.paragraph_format.left_indent = Cm(0.5)


def add_table(df, caption=None, float_fmt="{:.2f}", fontsize=9):
    if caption:
        p = doc.add_paragraph()
        r = p.add_run(caption)
        r.bold = True
        r.font.size = Pt(10)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t = doc.add_table(rows=df.shape[0] + 1, cols=df.shape[1] + 1)
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0].cells
    hdr[0].text = ""
    for j, c in enumerate(df.columns, start=1):
        hdr[j].text = str(c)
    for i in range(df.shape[0]):
        row = t.rows[i + 1].cells
        row[0].text = str(df.index[i])
        for j in range(df.shape[1]):
            v = df.iloc[i, j]
            row[j + 1].text = (float_fmt.format(v) if isinstance(v, (int, float, np.number)) else str(v))
    for row in t.rows:
        for cell in row.cells:
            for par in cell.paragraphs:
                for run in par.runs:
                    run.font.size = Pt(fontsize)
    doc.add_paragraph()
    return t


def add_figure(path, caption, width_cm=15):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(path, width=Cm(width_cm))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap.add_run(caption)
    r.italic = True
    r.font.size = Pt(9)


# ================================================================
# 封面
# ================================================================
add_heading("作业3：公平性实践", 0)
add_para("《负责任人工智能：原则、治理与量化方法》", align=WD_ALIGN_PARAGRAPH.CENTER)
add_para("黄斐，新南威尔士大学", align=WD_ALIGN_PARAGRAPH.CENTER)
add_para("姓名：林富强    专业：软件工程", align=WD_ALIGN_PARAGRAPH.CENTER)
doc.add_paragraph()

add_heading("概述与实现环境")
add_para("本作业在第三章两个案例研究（车险公平定价、COMPAS 审前风险评估）的基础上，完成六道题目："
         "A 部分车险定价的 MDP 的 lambda 扫描（问题1）、对另一个变量去偏（问题2）、为何仅靠无意识不够（问题3）"
         "与部署建议（问题4）；B 部分 COMPAS 分类的 theta 扩展扫描（问题5）与不存在满足所有准则的 theta（问题6）。")
add_para("实现采用 Python 3.12（Anaconda）：GLM 用 statsmodels（Poisson 频率 + Gamma 严重程度），"
         "梯度提升用 xgboost（count:poisson / reg:gamma，base_margin 传递风险敞口与预测频率），"
         "差异性影响消除器（disparate_impact_remover，Feldman et al. 2015）与拒绝选项修正（roc_pivot，"
         "Kamiran et al. 2012）严格对照 fairmodels/fairness 包源码在 Python 中复刻；分类公平性指标"
         "（人口统计学均等性、TPR 差距、预测率均等性）按 fairness 包定义计算。数据为 TASK3/pg15training_processed.csv"
         "（100,000 份保单）与 fairness 包 compas 数据集（两大族裔 5,278 名被告）。完整代码见 TASK3/py/ 目录。")

# ================================================================
# A 部分
# ================================================================
add_heading("A 部分——车险定价")

add_heading("基线复现：五种模型设计（GLM 与 XGBoost）", 2)
add_para("按讲义流程复现：M0（完整模型，含 Gender）、MU（无意识，移除 Gender）、MDP（所有预测变量去偏）、"
         "MCDP（仅对非合法变量 Insurancescore 去偏）、MC（拟合完整模型，评分时对性别取平均），"
         "每种设计分别以 GLM 与 XGBoost 拟合；频率—严重程度分解下纯保费 = 频率 × 严重程度 × 365，"
         "模型3-5 与全部 XGBoost 模型的组合保费总额均按比例缩放至与 GLM MU 一致。")

mt = mean_table.set_index(["Method", "Gender"])
glm_male = mt.loc[("GLM", "Male")]; glm_female = mt.loc[("GLM", "Female")]
xgb_male = mt.loc[("XGBoost", "Male")]; xgb_female = mt.loc[("XGBoost", "Female")]
tb = pd.DataFrame({
    "M0": [glm_male["M0: Full Model"], glm_female["M0: Full Model"], xgb_male["M0: Full Model"], xgb_female["M0: Full Model"]],
    "MU": [glm_male["MU: Unawareness Model"], glm_female["MU: Unawareness Model"], xgb_male["MU: Unawareness Model"], xgb_female["MU: Unawareness Model"]],
    "MDP": [glm_male["MDP: Demographic Parity"], glm_female["MDP: Demographic Parity"], xgb_male["MDP: Demographic Parity"], xgb_female["MDP: Demographic Parity"]],
    "MCDP": [glm_male["MCDP: Conditional Demographic Parity"], glm_female["MCDP: Conditional Demographic Parity"], xgb_male["MCDP: Conditional Demographic Parity"], xgb_female["MCDP: Conditional Demographic Parity"]],
    "MC": [glm_male["MC: Controlling for the Protected Variable"], glm_female["MC: Controlling for the Protected Variable"], xgb_male["MC: Controlling for the Protected Variable"], xgb_female["MC: Controlling for the Protected Variable"]],
}, index=["GLM male", "GLM female", "XGBoost male", "XGBoost female"])
add_table(tb, "表1  平均预测纯保费（按性别，与讲义表一致）", float_fmt="{:.2f}")
add_para("复现结果与讲义平均保费表吻合（GLM：male M0=130.47、MU=114.03；female M0=95.66、MU=124.05 等，"
         "MDP/MCDP 因去偏步骤的随机噪声（jitter）存在 ±0.3 的合理微差；XGBoost 差异小于 0.1）。")

fa_tb = fa.set_index(["Method", "Model"])[["male_mean", "female_mean", "disparate_impact_ratio", "rmse"]].round(5)
add_table(fa_tb, "表2  公平性—准确性汇总（DIR = male/female，RMSE 越小越好）", float_fmt="{:.5f}")
add_figure(os.path.join(OUT, "fig1_fairness_accuracy.png"),
           "图1   GLM 与 XGBoost 定价模型间的公平性—准确性比较")

# ---------------- 问题1 ----------------
add_heading("问题1 — MDP 的 lambda 扫描", 2)
add_para("（a）实现。lambda 是差异性影响消除器的去偏强度旋钮：lambda=0 精确等价于 MU（不去偏），"
         "lambda=1 等价于讲义 MDP，中间取值是讲义未计算的部分修正 MDP。对 lambda ∈ {0, 0.25, 0.5, 0.75, 1}，"
         "将全部六个预测变量（Age.ct、Bonus、Value、Density、GroupOne、Insurancescore）视为非合法变量去偏，"
         "重新拟合 GLM（Poisson 频率 + Gamma 严重程度），计算差异性影响比率（DIR = 男性平均保费 / 女性平均保费）"
         "与 RMSE。核心代码：")
add_code('''for lam in [0.0, 0.25, 0.5, 0.75, 1.0]:
    debiased = make_di_removed_data(ClaimsData, ALL_VARS, lambda_=lam)  # 全部变量去偏
    freq, sev = fit_glm_pair(debiased, debiased_rd, with_gender=False)  # 重拟合 GLM
    prem = predict_glm_premiums(debiased, freq, sev, with_gender=False) # 纯保费预测
    DIR = prem[male].mean() / prem[female].mean()
    RMSE = sqrt(mean((realclaim - prem)**2))''')
q1_tb = q1.set_index("lambda")
q1_tb.columns = ["DIR", "RMSE", "male_mean", "female_mean", "total"]
add_table(q1_tb[["DIR", "RMSE", "male_mean", "female_mean"]],
          "表3  问题1：五组 (lambda, DIR, RMSE) 结果", float_fmt="{:.4f}")
add_figure(os.path.join(OUT, "fig2_lambda_scan.png"),
           "图2   DIR 与 RMSE 随 lambda 变化的折线")
add_figure(os.path.join(OUT, "fig2b_lambda_fairness_acc.png"),
           "图3   公平性—准确性散点（每个点标注 lambda）", width_cm=12)
add_para("（c）解读。随着 lambda 从 0 增大到 1，DIR 从 0.938 单调上升至 1.007，越来越接近完全公平（DIR=1）；"
         "与此同时 RMSE 几乎水平（492.1–493.2，极差约 1），并未出现准确性显著恶化的现象。因此本数据上的规律是："
         "以可忽略的准确性代价换取公平性提升，而非标准的公平性—准确性权衡曲线。这与讲义的关键发现一致"
         "（公平性的代价比通常设想的要小）。两点补充：其一，lambda=0 的 DIR=0.938 与纯 MU（0.919）略有差异，"
         "原因是去偏流程自带的 jitter 噪声与分桶（pigeonholing）效应会轻微改变分布（fairmodels 文档亦明确说明）；"
         "其二，RMSE 的微小波动（约 ±1）远小于模型间差异量级，可视为数值噪声，结论稳健。")

# ---------------- 问题2 ----------------
add_heading("问题2 — 对另一个变量去偏", 2)
add_para("（a）实现方案。沿用 MCDP 的模式：仅对讲义从未去偏过的合法变量 Density 去偏（lambda=1），"
         "保持 Insurancescore 及其余所有预测变量不变，然后重新拟合 GLM 并计算 DIR 与 RMSE。"
         "在代码上只需把 make_di_removed_data() 的 features_to_transform 参数改为 [Density]。")
add_code('''deb_density = make_di_removed_data(ClaimsData, ["Density"], lambda_=1.0)
freq, sev   = fit_glm_pair(deb_density, deb_density_rd, with_gender=False)
prem        = predict_glm_premiums(deb_density, freq, sev, with_gender=False)
DIR, RMSE   = ...  # 同问题1''')
add_para("（b）结果。")
q2 = summary["q2_density_only"]
q2_tb = pd.DataFrame({
    "方案": ["MU（问题1 lambda=0）", "仅去偏 Density（lambda=1）", "MCDP（讲义，仅去偏 Insurancescore）", "MDP（lambda=1，全变量）"],
    "DIR": [q1.loc[q1["lambda"] == 0, "dir"].iloc[0], q2["dir"], 0.95695, q1.loc[q1["lambda"] == 1, "dir"].iloc[0]],
    "RMSE": [q1.loc[q1["lambda"] == 0, "rmse"].iloc[0], q2["rmse"], 492.99517, q1.loc[q1["lambda"] == 1, "rmse"].iloc[0]],
}).set_index("方案")
add_table(q2_tb, "表4  问题2：仅对 Density 去偏与各对照方案的 DIR/RMSE", float_fmt="{:.4f}")
add_figure(os.path.join(OUT, "fig4_debias_comparison.png"),
           "图4   各去偏方案的 DIR 对比")
add_para("（c）结果分析。仅对 Density 去偏后 DIR=0.931，与 MU（0.938）几乎相同，差异程度基本未被改变；"
         "RMSE 亦几乎不变（493.0 vs 492.4）。对照 MCDP（仅去偏 Insurancescore）DIR=0.957 更接近 1，"
         "全变量 MDP 则达到 1.007。原因结合第二章的代理变量概念：Density（居住密度）与受保护属性 Gender 的相关性很弱，"
         "并非性别信息的代理通道，去除其与 Gender 的关联对保费在两性间的再分配几乎没有影响；"
         "而 Insurancescore 作为信用评分代理（由类型、类别、职业、地区和年龄构建），与 Gender 高度相关，"
         "是性别差异进入模型的主要通道，因此对它去偏才能有效削弱代理性歧视。结论：去偏的效果取决于"
         "该变量与受保护属性的相关强度，而非机械地对任意变量去偏都有效。")

# ---------------- 问题3 ----------------
add_heading("问题3 — 为何仅靠无意识是不够的", 2)
add_para("（a）利用讲义平均保费表计算 M0 与 MU 的 DIR（= 男性平均保费 / 女性平均保费）：")
q3_tb = pd.DataFrame({
    "方法": ["GLM", "GLM", "XGBoost", "XGBoost"],
    "模型": ["M0", "MU", "M0", "MU"],
    "男性保费": [130.47, 114.03, 130.98, 114.41],
    "女性保费": [95.66, 124.05, 94.63, 123.39],
    "DIR": [q3["GLM_M0"]["DIR"], q3["GLM_MU"]["DIR"], q3["XGBoost_M0"]["DIR"], q3["XGBoost_MU"]["DIR"]],
}).set_index(["方法", "模型"])
add_table(q3_tb, "表5  讲义平均保费表下 M0/MU 的 DIR", float_fmt="{:.3f}")
add_para("观察：M0 中 DIR≈1.36–1.38（男性支付明显更多，超出五分之四原则区间 [0.8, 1.25]）；"
         "移除 Gender 得到 MU 后，DIR 显著下降至 0.92–0.93，明显更接近 1，但并未达到完全公平，"
         "且方向发生了反转——M0 中女性支付约少 27%，MU 中女性反而多支付约 8–9%。即谁付更多由男性变为女性。")
add_para("（b）结合第二章的代理性歧视概念：MU 虽然不再把 Gender 作为输入，但合法变量 Insurancescore"
         "（信用评分代理，由类型、类别、职业、地区和年龄构建）仍与 Gender 相关，充当了性别的代理变量，"
         "性别差异经由该代理变量重新进入定价模型。方向反转的原因是：M0 中 Gender 系数直接抵消了"
         "代理变量携带的性别效应；一旦移除 Gender，未被抵消的代理效应便主导结果，使女性整体被系统性地高估。"
         "这正说明无意识（把受保护属性移出输入集）本身很少足够：模型表面不含受保护属性，"
         "但预测仍系统性偏离均等性，法律上构成差异性影响（disparate impact）。")

# ---------------- 问题4 ----------------
add_heading("问题4 — 建议", 2)
add_para("基于问题1与问题2的结果，我建议部署 MCDP（条件人口统计学均等性，仅对非合法变量 Insurancescore 去偏，lambda=1）。")
add_para("论证：问题2表明去偏是否有效取决于变量是否为受保护属性的代理通道——对 Insurancescore 去偏（MCDP，DIR=0.957）"
         "能有效削弱歧视，而对与 Gender 无关的 Density 去偏则毫无效果（DIR=0.931，与 MU 相同）。"
         "MCDP 恰好只处理真正的代理通道（非合法变量 Insurancescore），保留年龄、Bonus、GroupOne、Density、Value"
         "等合法定价变量携带的风险区分信息，因此它在公平性（DIR=0.957 接近 1）与准确性（RMSE=493.0，"
         "与 MU 的 486.5 仅差约 1%）之间取得最佳平衡。这与第二章中条件人口统计学均等性的准则一致："
         "预测应在受保护属性上独立，但在合法特征（风险因子）条件下仍可存在合理差异，符合保险监管"
         "（如欧盟统一性别定价规则）认可的合法定价差异框架。相比之下，完全 MDP（lambda=1，全变量去偏）"
         "虽然把 DIR 推到 1.007、RMSE 代价同样很小（493.2），但它把合法风险变量也一并压平，"
         "破坏了定价模型的风险区分能力与可解释性，属于矫枉过正；完全不干预的 MU 则保留了代理性歧视"
         "（DIR=0.919，女性被系统高估约 9%），面临监管风险。若监管环境要求更强的人群均等性，"
         "亦可取中等 lambda（0.5–0.75）作为折中：DIR 达 0.98–0.99，RMSE 代价仍可忽略。")

# ================================================================
# B 部分
# ================================================================
add_heading("B 部分——COMPAS 分类")

add_heading("数据与模型复现", 2)
add_para(f"使用 fairness 包 compas 数据集，保留两大族裔后 N={compas['N']} 名被告"
         f"（白人 {compas['n_white']}，非裔美国人 {compas['n_afam']}）。"
         f"基础再犯率：白人 {compas['base_rate_white']*100:.1f}%，非裔美国人 {compas['base_rate_afam']*100:.1f}%——"
         f"这一基准率差异最终驱动了本部分几乎所有结果。预测变量为前科次数（Number_of_Priors，已标准化）、"
         f"两个年龄指标（Age_Above_FourtyFive / Age_Below_TwentyFive）、性别（Female）、轻罪标志（Misdemeanor）。"
         f"M0 额外加入族裔（参考水平 Caucasian），MU 移除族裔；逻辑回归以 0.5 为分界，"
         f"预测再犯概率高于 0.5 即标记为高风险。")
m0 = compas["m0"]; mu = compas["mu"]
compas_base_tb = pd.DataFrame({
    "指标": ["准确率", "被标记高风险：白人", "被标记高风险：非裔", "FPR：白人/非裔", "FNR：白人/非裔", "精确率：白人/非裔",
             "人口统计学均等性", "机会均等（TPR 差距）", "预测率均等性"],
    "M0（含族裔）": [f"{m0['accuracy']*100:.1f}%", f"{m0['flagged_white']*100:.1f}%", f"{m0['flagged_African_American']*100:.1f}%",
                     f"{m0['fpr_white']*100:.1f}% / {m0['fpr_African_American']*100:.1f}%",
                     f"{m0['fnr_white']*100:.1f}% / {m0['fnr_African_American']*100:.1f}%",
                     f"{m0['precision_white']*100:.1f}% / {m0['precision_African_American']*100:.1f}%",
                     f"{m0['prop_parity']:.2f}", f"{m0['equal_odds']:.2f}", f"{m0['pred_rate_parity']:.2f}"],
    "MU（无意识）": [f"{mu['accuracy']*100:.1f}%", f"{mu['flagged_white']*100:.1f}%", f"{mu['flagged_African_American']*100:.1f}%",
                     f"{mu['fpr_white']*100:.1f}% / {mu['fpr_African_American']*100:.1f}%",
                     f"{mu['fnr_white']*100:.1f}% / {mu['fnr_African_American']*100:.1f}%",
                     f"{mu['precision_white']*100:.1f}% / {mu['precision_African_American']*100:.1f}%",
                     f"{mu['prop_parity']:.2f}", f"{mu['equal_odds']:.2f}", f"{mu['pred_rate_parity']:.2f}"],
}).set_index("指标")
add_table(compas_base_tb, "表6  M0 / MU 的分类指标与三大准则比率（与讲义表一致）", float_fmt="{}")
add_para("复现结果与讲义完全一致：无意识几乎没有改变任何东西——MU 的准确率仅下降 0.3 个百分点"
         "（66.5% vs 66.8%），但所有差异比率几乎原封不动（人口统计学均等性 1.94，TPR 差距 1.67，"
         "预测率均等性 1.15），一名不会再犯的非裔被告被标记高风险的概率约为不会再犯白人的两倍"
         "（FPR：33.0% vs 17.2%），复现了 ProPublica 对真实 COMPAS 工具的核心发现。")

# ---------------- 问题5 ----------------
add_heading("问题5 — 扩展 theta 扫描", 2)
add_para("（a）实现方案。roc_pivot（拒绝选项修正）对落在决策分界值附近、宽度为 theta 的区间内的预测做"
         "镜像翻转：优势群体（白人）中略高于分界值（被预测有利=不再犯）的个案翻转为不利，弱势群体（非裔）中"
         "略低于分界值的个案翻转为有利；区间之外的有把握预测保持不变。由于本模型中高于 0.5 意味着预测再犯"
         "（不利结果），先将模型转换为预测不再犯这一有利结果，保持白人（现状模型所偏向的群体）为优势群体，"
         "再调用 roc_pivot。对 MU 在 theta=0.25 与 0.30 下应用，并沿用讲义对 M0/MU 所用的同一套指标"
         "（准确率、人口统计学均等性、TPR 差距、预测率均等性）。核心代码：")
add_code('''p_norecid = 1.0 - p_recid_mu                     # 转换为预测不再犯（有利结果）
p_fixed   = roc_pivot(p_norecid, ethnicity,          # 拒绝选项修正
                      privileged="Caucasian", cutoff=0.5, theta=theta)
pred      = np.where(p_fixed > 0.5, "no", "yes")     # 修正后的再犯标签
# 指标：accuracy / prop_parity / equal_odds / pred_rate_parity（fairness 包同款定义）''')
add_para("（b）结果。")
scan_tb = scan.copy()
scan_tb.columns = ["theta", "准确率", "人口统计学均等性", "机会均等(TPR 差距)", "预测率均等性", "被标高风险：白人", "被标高风险：非裔"]
scan_tb["准确率"] = (scan_tb["准确率"] * 100).round(1).astype(str) + "%"
scan_tb = scan_tb.drop(columns=["被标高风险：白人", "被标高风险：非裔"])
scan_tb = scan_tb.set_index("theta")
add_table(scan_tb, "表7  roc_pivot theta 扫描（0 至 0.30，0–0.20 与讲义一致，0.25/0.30 为本作业新增）",
          float_fmt="{:.3f}")
add_figure(os.path.join(OUT, "fig3_compas_theta_scan.png"),
           "图5   COMPAS roc_pivot theta 扫描：四大指标随 theta 变化")
add_para("（c）结果分析。讲义中的模式在更高 theta 下持续并加剧：人口统计学均等性进一步超调至 1 以下"
         "（theta=0.25 时 0.162，theta=0.30 时 0.100），TPR 差距同样继续下冲（0.217 → 0.149），"
         "预测率均等性持续恶化（1.789 → 1.982），准确率不断下降（54.9% → 50.4%，接近随机猜测）。"
         "存在明确的临界点：theta≈0.10–0.15 之后，修正已从消除差异变成反转差异"
         "（人口统计学均等性下冲至 1 以下意味着非裔反而被标记高风险的比率远低于白人）；"
         "theta>0.20 后模型失去实际意义——准确率逼近 50%，而被标记高风险的非裔占比仅剩白人的约十分之一，"
         "属于系统性反向歧视。")

# ---------------- 问题6 ----------------
add_heading("问题6 — 为何不存在满足所有准则的 theta", 2)
add_para("第二章的不可能性结论指出：当两个群体的基础结果率（基准率）不同时，"
         "人口统计学均等性（预测独立于受保护属性）、均衡赔率/分离性（在预测条件下结果独立于受保护属性）"
         "与预测率均等性/充分性（在结果条件下预测独立于受保护属性）无法同时满足，除非预测完美。"
         "COMPAS 数据中白人再犯率 39.1%、非裔 52.3%，两者不同，因此不存在能够同时使三项比率都接近 1 的单一 theta 值。")
add_para("roc_pivot 只能沿着这条不可能曲线移动：它把人口统计学均等性与 TPR 差距推向 1 的同时，"
         "必然以牺牲预测率均等性为代价（表7 中预测率均等性从 1.15 单调恶化到 1.98）。这是因为三类准则"
         "以不同的方式处理基准率差异——人口统计学均等性完全忽略真实结果，因而同时反映模型差异与基准率差异；"
         "精确率以预测为条件，部分吸收了基准率差异。当基准率不同时，把其中两者推向均等必然使第三者失衡。")
add_para("要使这样的 theta 存在，需要两个群体满足的条件是：基础再犯率相同"
         "（即白人与非裔的再犯率相等，此时三项准则可由同一阈值同时满足），"
         "或者模型在受保护属性上达到完美预测。现实中这两个条件均不成立，"
         "因此选择一种后处理修正，本质上就是选择将基准率差异中的哪一部分计入不公平、"
         "哪一部分视为信号——这与第二章第一步选择准则密不可分。")

# ================================================================
# 附录
# ================================================================
add_heading("附录：代码清单", 2)
add_para("全部代码位于 TASK3/py/ 目录，可直接运行复现：")
add_code('''TASK3/py/
├── fairness_utils.py      # DI remover / roc_pivot / 分类公平性指标（对照 R 源码复刻）
├── insurance_models.py    # 数据加载、GLM/XGBoost 拟合、预测与指标
├── 00_baseline.py         # 基线复现：5 模型 x (GLM|XGBoost)，输出表1/表2/图1
├── 01_q1_q2.py            # 问题1 lambda 扫描 + 问题2 Density 去偏
├── 02_q3.py               # 问题3 DIR 计算（讲义平均保费表）
├── 03_plots.py            # 全部图表的生成
├── 05_q5_compas.py        # COMPAS M0/MU + theta 扩展扫描（表6/表7/图5）
└── output/                # 结果表（csv/json）与图表（png）''')

# 保存
docx_path = os.path.join(TASK3, "林富强_作业3.docx")
doc.save(docx_path)
print("已生成:", docx_path)
