# -*- coding: utf-8 -*-
# ============================================================
# build_作业7.py — 生成《林富强_作业7.docx》并通过 LibreOffice 转为 PDF
# 内容：隐私实践 — 三道题的代码、输出、图表与书面回答
# 运行：python build_作业7.py   （依赖 python-docx；先运行 01–04 分析脚本）
# ============================================================
import os
import subprocess
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
OUT = os.path.join(BASE, "output")

doc = Document()
style = doc.styles["Normal"]
style.font.name = "Times New Roman"
style.font.size = Pt(11)
style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
for sec in doc.sections:
    sec.top_margin = Cm(2.2)
    sec.bottom_margin = Cm(2.2)
    sec.left_margin = Cm(2.5)
    sec.right_margin = Cm(2.5)


def _zh(run, name, size):
    run.font.name = "Times New Roman"
    rPr = run._element.get_or_add_rPr()
    rF = rPr.find(qn("w:rFonts"))
    if rF is None:
        rF = OxmlElement("w:rFonts"); rPr.append(rF)
    rF.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)


def H(text, level=1):
    h = doc.add_heading(text, level=level)
    for r in h.runs:
        _zh(r, "黑体", {0: 17, 1: 14, 2: 12}.get(level, 12))
        r.font.color.rgb = RGBColor(0, 0, 0)
    return h


_BOLD = None


def P(text, bold=False, size=11, align=None, indent=True):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    if indent:
        p.paragraph_format.first_line_indent = Cm(0.74)
    if align:
        p.alignment = align
    # 支持 **bold** 标记
    import re
    pos = 0
    for m in re.finditer(r"\*\*(.+?)\*\*", text):
        if m.start() > pos:
            r = p.add_run(text[pos:m.start()]); _zh(r, "宋体", size); r.bold = bold
        r = p.add_run(m.group(1)); _zh(r, "宋体", size); r.bold = True
        pos = m.end()
    if pos < len(text):
        r = p.add_run(text[pos:]); _zh(r, "宋体", size); r.bold = bold
    return p


def CODE(text, size=8.5):
    lines = text.rstrip("\n").split("\n")
    for ln in lines:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.left_indent = Cm(0.5)
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear"); shd.set(qn("w:fill"), "F2F2F2")
        p._p.get_or_add_pPr().append(shd)
        r = p.add_run(ln if ln else " ")
        r.font.name = "Consolas"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        r.font.size = Pt(size)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def TBL(header, rows, widths=None, fs=9):
    t = doc.add_table(rows=len(rows) + 1, cols=len(header))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, h in enumerate(header):
        c = t.rows[0].cells[j]; c.text = ""
        r = c.paragraphs[0].add_run(h); _zh(r, "黑体", fs); r.bold = True
        c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    for i, row in enumerate(rows):
        for j, v in enumerate(row):
            c = t.rows[i + 1].cells[j]; c.text = ""
            r = c.paragraphs[0].add_run(str(v)); _zh(r, "宋体", fs)
            c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    if widths:
        for row in t.rows:
            for j, w in enumerate(widths):
                row.cells[j].width = Cm(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t


def IMG(name, width=6.1, caption=None):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(os.path.join(OUT, name), width=Inches(width))
    if caption:
        cp = doc.add_paragraph(); cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = cp.add_run(caption); _zh(r, "楷体", 9)


# ================================================================
# 封面
# ================================================================
H("作业7：隐私实践", 0)
P("《负责任人工智能：原则、治理与量化方法》", align=WD_ALIGN_PARAGRAPH.CENTER, indent=False)
P("黄斐，新南威尔士大学", align=WD_ALIGN_PARAGRAPH.CENTER, indent=False)
P("姓名：林富强    专业：软件工程", align=WD_ALIGN_PARAGRAPH.CENTER, indent=False)
doc.add_paragraph()

# ================================================================
# 概述
# ================================================================
H("概述", 1)
P("本作业在第七章车险案例研究（pg15training，100,000 份法国第三方责任险保单）的基础上，"
  "把「识别准标识符 → 选择隐私增强技术（PET）→ 应用 → 衡量权衡取舍 → 决策」的完整工作流程，"
  "应用到讲义尚未计算过的具体参数上：移除 AgeGroup 后的重新识别风险评估与 k=100 的 k-匿名化（问题 1）、"
  "针对 66 岁以上密集城区子群体的差分隐私扫描（问题 2）、以及把 TSTR 评估拓展到预测 Gender 的替代模型（问题 3）。")
P("**数据与预处理：**复用讲义的 data/pg15training_raw.csv（剥离直接标识符后的 8 列：Gender、Age、AgeGroup、"
  "DensityGroup、ValueGroup、BonusGroup、HasClaim、ClaimAmount）。准标识符分组与讲义逐条核对一致——前 6 条记录的"
  "分组结果与讲义示例表完全相同，且关键锚点全部复现：女性 36,568 人 / 男性 63,432 人、平均年龄 41.13、"
  "理赔率 0.1226、男性 66 岁以上群体 5,031 条。其中 ValueGroup、DensityGroup、BonusGroup 的具体切点讲义未公开，"
  "本文按与示例记录一致的原则自行标定（Value：5000/12000/28000/45000；Density：50/100/200；Bonus：-25/25/75），"
  "因此五变量基线（全局风险 0.943%）与讲义（0.77%）存在约 0.2 个百分点的差异；"
  "作业要求的所有对比均在**同一套分组**内部进行，结论不受该差异影响。")
P("**实现说明：**本作业原题为 R 实践。因本机未安装 sdcMicro/synthpop，以下分析用 Python 对讲义工作流做等价复现："
  "sdcMicro 的个体风险 r_k = 1/f_k（文件按总体处理）、局部抑制的最少单元格贪心算法与 NA 通配匹配语义、"
  "拉普拉斯机制（自实现）、以及 synthpop 的 CART 序列合成（minbucket=5，叶内均匀抽取真实观测值），"
  "方法与讲义一一对应。全部代码、输出与图表如下。")

# ================================================================
# 问题 1
# ================================================================
H("问题 1：移除 AgeGroup 后评估重新识别风险并应用 k-匿名性", 1)

H("关键代码", 2)
CODE("""# 准标识符：讲义五变量 vs 移除 AgeGroup 后的四变量
QI5 = ["Gender", "AgeGroup", "DensityGroup", "ValueGroup", "BonusGroup"]
QI4 = ["Gender", "DensityGroup", "ValueGroup", "BonusGroup"]   # 问题1的 keyVars

# ---- createSdcObj 等价实现：个体风险 r_k = 1/f_k（总体语义）----
def risk_summary(df, key_vars):
    key = df[key_vars].astype(str)
    fk = key.groupby(key_vars).transform("size").to_numpy()
    rk = 1.0 / fk
    return {"global_risk_pct": 100*rk.sum()/len(df),      # = 预期重识别数/n
            "expected_reid": rk.sum(), "max_risk": rk.max(),
            "k1": (fk==1).sum(), "k3": (fk<=3).sum(), "k5": (fk<=5).sum()}

base5, base4 = risk_summary(df_raw, QI5), risk_summary(df_raw, QI4)""")

CODE("""# ---- k=100 局部抑制（localSuppression 等价实现）----
# importance = c(1,2,3,4)：Gender 最受保护，BonusGroup 最可舍弃
# 抑制子集按（个数升序，importance 降序）枚举，NA 按通配符匹配
K, IMPORTANCE = 100, {"Gender":1, "DensityGroup":2, "ValueGroup":3, "BonusGroup":4}
for i in viol_idx:                       # 等价类 < K 的记录
    for T in subsets:                    # T = 拟抑制的变量集合
        S = tuple(v for v in QI4 if v not in T)
        if size_lookup[S][i] >= K:       # 通配符等价类达到 K 即停止
            for v in T: suppressed[v][i] = True
            break""")

H("a. 四变量（移除 AgeGroup）的重新识别风险", 2)
P("按讲义表格样式汇报。基线为同一数据、同一管线下的五变量结果（讲义五变量结果附于括号内以便对照）：")
TBL(
    ["指标", "5变量基线（讲义）", "4变量：移除 AgeGroup"],
    [
        ["全局重新识别风险", "0.943%（讲义 0.765%）", "0.160%"],
        ["预期重新识别数量", "943（讲义 765）", "160"],
        ["最大个体风险", "1.00", "0.10"],
        ["k = 1 的记录数", "27（讲义 6）", "0"],
        ["k ≤ 3 的记录数", "126（讲义 40）", "0"],
        ["k ≤ 5 的记录数", "350（讲义 102）", "0"],
    ],
    widths=[5.2, 5.6, 5.2],
)
P("移除 AgeGroup 后，四变量的全部可能组合只有 2×4×5×4 = 160 种，实际恰好占满 160 种组合，"
  "因此全局风险 = 160/100,000 = **0.160%**；数据集中**不再存在任何等价类规模 ≤ 5 的记录**，"
  "最小的等价类也有 10 条记录（女性、密集城区、Very high 车价、High 折扣档），最大个体风险从 1（唯一记录）降至 0.1。")

H("b. 与五变量基准的比较：AgeGroup 是否「最具识别性」？", 2)
P("**风险大幅下降**：全局风险从 0.943% 降至 0.160%（相对降幅约 83%），预期重新识别数量从 943 条降至 160 条；"
  "更关键的是结构变化——所有 27 条唯一记录和全部 350 条 k ≤ 5 的记录**完全消失**。")
P("这一变化**与讲义「AgeGroup（或 Age）是最具识别性的变量之一」的说法高度相符**：五个变量中 Gender 只有 2 个水平、"
  "DensityGroup/ValueGroup/BonusGroup 各 4–5 个水平，而 6 档年龄段与它们的叉积提供了最细的切分粒度——"
  "943 种被占用的五变量组合中，绝大多数正是在 AgeGroup 这一维上才彼此区分。把这一维整体移除后，"
  "原来仅靠年龄段差异而彼此孤立的记录全部并入同一等价类，风险随之坍缩。换一个角度说：讲义五变量 k=50 抑制后风险仍剩 0.41%，"
  "而这里仅仅删掉一列（不做任何抑制）就把风险压到 0.16%——**AgeGroup 承载了该数据集大部分的重识别风险**。"
  "同时也应看到反面：0.16% 并不为零，160 个等价类中仍有一批 10–32 条的小类，说明「删除最识别性变量」不等于匿名。")

H("c. k = 100 局部抑制的结果", 2)
P("在同一四变量集合上以 importance = c(1,2,3,4)（Gender → DensityGroup → ValueGroup → BonusGroup，与讲义同序）应用 k=100 "
  "局部抑制。需要处理的记录（等价类 < 100）共 1,767 条，抑制结果如下：")
TBL(
    ["准标识符", "被抑制取值个数", "抑制率"],
    [
        ["Gender", "0", "0.00%"],
        ["DensityGroup", "0", "0.00%"],
        ["ValueGroup", "0", "0.00%"],
        ["BonusGroup", "1,767", "1.77%"],
    ],
    widths=[5.2, 5.6, 5.2],
)
P("被抑制单元格总计 1,767 个。抑制后**全局风险降至 0.124%**（预期重新识别数 124 条），最大个体风险恰为 0.010 = 1/k，"
  "即每条记录的通配符等价类都至少包含 k = 100 条记录——k-匿名性保证严格成立。")
IMG("fig_q1_suppression_rate.png", 5.6, "图 1：k=100、importance 降序牺牲下各准标识符的抑制率")
IMG("fig_q1_bonus_dist.png", 6.1, "图 2：抑制前后 BonusGroup 分布。1,767 条记录（1.77%）的折扣档被置为 NA")

H("d. 与讲义 k=50、五变量结果的比较", 2)
P("**抑制不但没有扩散，反而更集中了。**讲义在 k=50、五变量下抑制了 7,339 个单元格（BonusGroup 6.90% + ValueGroup 0.44%，"
  "94% 的抑制落在 BonusGroup，其余溢出到了 ValueGroup），全局风险 0.77% → 0.41%；本题在**翻倍的 k=100**、四变量下"
  "却只抑制了 1,767 个单元格（1.77%），且**全部**落在 BonusGroup，Gender、DensityGroup、ValueGroup 一个都没动，"
  "最终风险 0.124% 还远低于讲义的 0.41%。")
P("**用 importance 排序解释：**localSuppression 总是先耗尽代价最低（importance 数值最大）的变量，只有当仅抑制该变量"
  "不足以使某条记录的等价类达到 k 时，才会动用下一个受保护程度更高的变量。讲义的五变量集合中存在大量仅靠抑制 BonusGroup "
  "无法修复的微小等价类——例如「女性、66+、密集城区、豪华车」这类 (Gender, AgeGroup, DensityGroup, ValueGroup) 组合"
  "本身已不足 50 条，必须继续牺牲 ValueGroup 才能达标，于是出现 438 个 ValueGroup 抑制。而移除 AgeGroup 后，"
  "最小的三变量组合 (Gender, DensityGroup, ValueGroup) 也有 166 条记录 ≥ k=100：**任何记录只需抑制一个 BonusGroup 取值**"
  "即可并入规模 ≥ 166 的通配符等价类，ValueGroup 永远轮不到被抑制。")
P("**预先移除 AgeGroup 是否改变这一比较？——是，而且是质变。**AgeGroup 是小等价类的最主要来源；预先把它从 keyVars 中拿掉，"
  "相当于在抑制开始之前就把最细的切分维度抹平，于是「k 翻倍、抑制反而缩减为 1/4」这一反直觉结果得以出现。"
  "这也印证了讲义的提醒：抑制率随 k 加速上升的前提是准标识符集合本身能切出大量小类——切分维度减少后，"
  "同一个 k 对应的效用代价可能不升反降。")

# ================================================================
# 问题 2
# ================================================================
H("问题 2：针对 66 岁以上、密集城区保单持有人的差分隐私", 1)
H("关键代码", 2)
CODE("""# 拉普拉斯机制：M(D) = f(D) + Lap(Δf/ε)，Age 截断到 [18,100]
def dp_mean(x, epsilon, lower=18, upper=100):
    x = np.clip(x, lower, upper); n = len(x)
    sens = (upper - lower) / n              # Δf = (upper-lower)/n
    u = rng.uniform(-0.5, 0.5)
    return x.mean() - sens/epsilon * np.sign(u) * np.log(1 - 2*abs(u))

sub = df_raw[(df_raw.AgeGroup == "66+") & (df_raw.DensityGroup == "Dense urban")]
EPS = [0.1, 0.5, 1.0, 5.0, 10.0]
for eps in EPS:                            # 每个 ε 独立重复发布 200 次
    est = [dp_mean(sub.Age, eps) for _ in range(200)]""")
P("（同样的扫描也运行于整体业务组合 n=100,000 与讲义基准的男性 66+ 豪华车子群体，作为 2c 比较的两组锚点。）", indent=False, size=10)

H("a. 子群体规模与扫描设置", 2)
P("筛选 AgeGroup == \"66+\" 且 DensityGroup == \"Dense urban\" 的子群体，规模 **n = 1,490**，"
  "真实平均年龄 70.26 岁。对该子群体的 Age 变量在讲义同样的五个隐私预算 ε ∈ {0.1, 0.5, 1, 5, 10} 下，"
  "各独立重复发布 200 次。")

H("b. 各 ε 下的噪声模式", 2)
TBL(
    ["ε", "噪声 SD（岁）", "平均 |偏差|（岁）", "最大 |偏差|（岁）", "实践解读"],
    [
        ["0.1", "0.840", "0.577", "5.14", "单次发布可偏离真实均值 5 岁以上"],
        ["0.5", "0.165", "0.122", "0.57", "偏差通常小于 0.6 岁，尾部仍可达 1 岁以上"],
        ["1.0", "0.066", "0.048", "0.22", "常用默认值，精度已可接受"],
        ["5.0", "0.016", "0.011", "0.07", "即便对该子群体也几乎无感"],
        ["10.0", "0.008", "0.005", "0.03", "接近未加噪输出"],
    ],
    widths=[1.6, 3.2, 3.4, 3.4, 4.4],
    fs=9.5,
)
IMG("fig_q2_dp_box.png", 6.3, "图 3：200 次重复发布中 DP 平均年龄估计偏差的箱型图（同一机制、同一 ε，仅 n 不同）")
P("图 3 中，整体业务组合（蓝）在所有 ε 下都紧贴零线；本题的 66+ 密集城区子群体（绿）与讲义的男性 66+ 豪华车"
  "基准（橙，n=548）则展现出同一现象的两个强度：ε=0.1 时前者单次发布可偏离约 5 岁、后者可超过 6 岁；"
  "到 ε ≥ 5 时两者箱体才收窄到 0.2 岁以内。**机制与隐私预算完全相同，唯一的差别是 n。**")

H("c. 与敏感度公式 (upper−lower)/n 的定量核对", 2)
P("三个群体的敏感度与 ε=0.1 时的理论噪声尺度（拉普拉斯分布 SD = √2·Δf/ε）如下：")
TBL(
    ["群体", "n", "Δf = 82/n", "理论 SD（ε=0.1）", "实测 SD（ε=0.1）"],
    [
        ["整体业务组合", "100,000", "0.00082", "0.012", "0.010"],
        ["本题：66+ 且密集城区", "1,490", "0.05503", "0.778", "0.840"],
        ["讲义基准：男性66+豪华车", "548", "0.14964", "2.116", "2.011"],
    ],
    widths=[5.4, 2.2, 2.6, 3.2, 2.8],
)
P("**观察到的噪声量与敏感度公式的预测相符。**三个群体使用完全相同的机制与 ε，噪声 SD 按 1/n 缩放："
  "本题子群体与讲义基准子群体的实测 SD 之比为 2.011/0.840 ≈ 2.4，与两群体规模之比的倒数 1,490/548 ≈ 2.7 同量级"
  "（200 次重复的蒙特卡洛误差范围内）；本题子群体与整体组合的敏感度之比为 0.055/0.00082 ≈ 67，"
  "对应理论噪声 SD 之比同样约 67 倍（0.778 对 0.012），实测 0.840 对 0.010 与之相符。"
  "换算成直觉：ε=0.1 时在 n=1,490 的子群体上，**一次发布的平均年龄可以错 5 岁**，而这已是隐私保护的「免费午餐」"
  "用完之后的代价——若监管要求按更细的口径（如单一邮编）分解，n 再缩一个量级，同样的 ε 就会带来完全不可用的输出，"
  "必须提高 ε、放粗粒度或改用其他 PET。这也与讲义结论一致：选择 ε 与决定发布口径的精细程度无法分开考虑。")

# ================================================================
# 问题 3
# ================================================================
H("问题 3：合成数据——用于预测 Gender 的 TSTR", 1)
H("关键代码", 2)
CODE("""# ---- synthpop(CART) 式序列合成（minbucket=5，叶内均匀抽取真实值）----
VISIT = ["GenderMale", "Age", "Density", "Value", "Bonus", "HasClaim", "ClaimAmount"]
for var in VISIT:
    preds = [p for p in VISIT if p in syn.columns]
    tree = DecisionTree(Classifier/Regressor)(min_samples_leaf=5).fit(real[preds], real[var])
    leaf = tree.apply(syn[preds])            # 合成记录走到叶节点
    syn[var] = [random.choice(叶内真实取值) for lf in leaf]
# df_synth_input (80,000) / df_real_holdout (2,000)：直接复用讲义划分

# ---- 问题3模型：预测 Gender（只换公式与目标）----
X = ["Age", "log1p(Density)", "log1p(Value)", "Bonus"]
m_real = LogisticRegression().fit(df_synth_input[X], df_synth_input.GenderMale)
m_syn  = LogisticRegression().fit(syn_df[X],         syn_df.GenderMale)
auc_manual(y_holdout, m_real.predict_proba(X_holdout)[:,1])  # 秩(Mann-Whitney) AUC""")

H("a. 模型设定", 2)
P("按题目要求，仅改变模型公式与预测目标：以 **Age、log1p(Density)、log1p(Value)、Bonus** 四个预测因子拟合逻辑回归预测 "
  "**Gender**，分别在 df_synth_input（真实数据，80,000 条）与 syn_df（CART 合成数据，80,000 条）上各训练一次，"
  "并在讲义构建的 2,000 条真实留出记录（合成过程从未见过）上以 auc_manual() 评估。留出集中男性占 65.0%。")

H("b. TSTR 结果", 2)
TBL(
    ["Model trained on", "AUC on real holdout (predicting Gender)"],
    [
        ["Real data", "0.5386"],
        ["Synthetic data", "0.5387"],
    ],
    widths=[7.0, 7.0],
)
P("合成数据训练的模型（0.5387）与真实数据训练的模型（0.5386）在真实留出集上的 AUC 几乎重合（差距 0.0001，甚至方向为正），"
  "与讲义理赔频率模型 0.728 对 0.729 的「差距可忽略」模式相同。")

H("c. 与 0.5 比较、与讲义差距比较及 CART 层面的解释", 2)
P("**第一步：AUC 与 0.5 的比较——这些变量对 Gender 几乎没有预测能力。**两个模型的 AUC 都只有约 0.54，"
  "仅比毫无预测能力的 0.5 高出 0.04：无论在真实数据还是合成数据上，Age、Density、Value、Bonus 四个变量合在一起"
  "也只包含关于持有人性别的微弱信号。这与数据直觉一致——车险定价因子中没有一个强性别标记变量，"
  "性别在理赔频率 GLM 中的系数（男性更高）反映的是「给定这些因子下男女理赔行为不同」，而非「这些因子能反推性别」。")
P("**第二步：两个 AUC 的差距是否比讲义更大？——不，反而更小（0.0001 对讲义的 ≤0.001）。**但第一步的答案确实"
  "改变了这一差距的含义：当真实信号本身就微弱到把 AUC 压在 0.54 附近时，TSTR 差距接近零只是「地板效应」——"
  "一个即使把性别相关结构扭曲得很厉害的合成数据集，也几乎不可能在一个本来就近乎随机的预测任务上显出性能落差。"
  "**近零的差距在这里不能作为「合成数据质量好」的有力证据**，它的信息量远低于讲义在强信号任务（理赔频率 AUC≈0.73）上"
  "得到的近零差距。")
P("**合成数据在这里确实更吃力，且在别处留下了痕迹：**对同一批数据拟合理赔频率 GLM 可以看到，真实数据上 GenderMale "
  "系数为 0.333，而合成数据训练的同一模型系数衰减为 0.261（差 -0.072，约 22%），与讲义观察到的方向一致"
  "（0.404 → 0.328）；其余变量（Age、log1p(Density)、log1p(Value)、Bonus）的系数则几乎无损"
  "（差距 0.000–0.011）。**用 synthpop 的 CART 生成方式可以给出一个理由：**CART 逐变量生成——Gender 位于访问序列最前端，"
  "其与后续变量的联合结构完全依赖「以已合成变量为条件的 CART 划分」来传递；树的每个叶节点内以常数化的经验分布抽样"
  "（minbucket=5），只能保留粗粒度的条件差异，性别在 Age/Value/Bonus 条件分布中的细弱结构在逐层条件化中被"
  "平滑衰减——强信号（Bonus、Age 对理赔的影响）经叶平均后存活，而性别这类弱条件信号则被系统性削弱。"
  "因此：预测理赔频率时合成数据几乎无损失，而任何依赖「性别 × 其他变量」细弱条件结构的任务（如本例的性别反推）"
  "都会首当其冲——只是该削弱恰好被 AUC≈0.5 的地板效应掩盖，需要用系数一致性这类检验才能暴露。")
IMG("fig_q3_marginals.png", 6.3, "图 4：合成数据的边际保真度——Age（原尺度）与 Density、Value（log1p 尺度）的真实/合成分布几乎重合")
P("补充的保真度检查与讲义结论一致：均值年龄 41.13 对 41.13、理赔频率 12.2% 对 12.5%、女性占比 36.7% 对 36.7%、"
  "车辆价值中位数 14,600 对 14,605——CART 合成在边际层面几乎无损，其隐私代价主要藏在联合结构（如上文的性别条件结构）中，"
  "完整的评估还应辅以成员推断与属性推断检验（Jordon 等 2022）。")

# ================================================================
# 总结
# ================================================================
H("总结", 1)
for txt in [
    "**重新识别风险由最细的准标识符主导。**移除单一最具识别性的 AgeGroup 使全局风险坍缩 83%（0.943%→0.160%）、"
    "全部小等价类消失；同一数据上「k 翻倍但抑制缩减为 1/4 且不溢出」的反直觉结果，均源于切分粒度的变化。",
    "**importance 排序决定了抑制的落点。**局部抑制总是先耗尽最可舍弃的变量（BonusGroup），"
    "只有当残留等价类小到抑制它也不够时才波及受保护变量；预先移除 AgeGroup 后最小的三变量组合仍有 166 条，"
    "使 k=100 的抑制 100% 集中在 BonusGroup（1.77%），Gender/Density/Value 零损失。",
    "**差分隐私的噪声按 1/n 缩放。**n=1,490 的 66+ 密集城区子群体在 ε=0.1 时单次发布可偏离 5 岁，"
    "而整体组合层面噪声仍可忽略；实测噪声 SD 与 √2·(82/n)/ε 的理论预测一致——选 ε 必须同时选定发布口径。",
    "**合成数据的 TSTR 表现取决于下游信号的强弱。**性别反推任务本身近随机（AUC≈0.54），近零差距是地板效应而非质量证明；"
    "性别条件结构在 CART 序列合成中被衰减约 22%（GenderMale 系数 0.333→0.261），评估合成数据应组合多种检验。",
]:
    P(txt)
P("所有数字均由随文脚本（01_prepare.py、02_q1_risk.py、03_q2_dp.py、04_q3_tstr.py）在 pg15training "
  "100,000 条真实保单上计算得到，随机种子固定，结果可复现。", indent=False, size=10)

# ---------- 保存并转 PDF ----------
out_docx = os.path.join(BASE, "林富强_作业7.docx")
doc.save(out_docx)
print("saved:", out_docx)
try:
    subprocess.run([r"C:\Program Files\LibreOffice\program\soffice.exe", "--headless",
                    "--convert-to", "pdf", out_docx, "--outdir", BASE], check=True, timeout=300)
    print("pdf converted")
except Exception as e:
    print("PDF 转换失败（可手动用 Word/WPS 另存）:", e)
