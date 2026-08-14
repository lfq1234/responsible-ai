# -*- coding: utf-8 -*-
"""生成 林富强_作业2（公平性原则）docx"""
import docx
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

doc = docx.Document()

# ---------- 全局样式 ----------
style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(11)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

def set_cn_font(run, name='宋体', size=None, bold=None, color=None):
    run.font.name = 'Times New Roman'
    r = run._element.rPr.rFonts
    r.set(qn('w:eastAsia'), name)
    if size: run.font.size = size
    if bold is not None: run.font.bold = bold
    if color: run.font.color.rgb = color

def add_heading(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    set_cn_font(run, name='黑体', size=Pt(14), bold=True)
    return p

def add_q(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    set_cn_font(run, name='宋体', size=Pt(11), bold=True)
    return p

def add_body(text, bold=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.3
    run = p.add_run(text)
    set_cn_font(run, name='宋体', size=Pt(11), bold=bold)
    return p

def add_body_mixed(parts):
    """parts: list of (text, bold)"""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.3
    for text, bold in parts:
        run = p.add_run(text)
        set_cn_font(run, name='宋体', size=Pt(11), bold=bold)
    return p

# ---------- 标题 ----------
t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = t.add_run('负责任人工智能 · 第二次作业：公平性原则')
set_cn_font(run, name='黑体', size=Pt(16), bold=True)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = sub.add_run('姓名：林富强　　专业：软件工程')
set_cn_font(run, name='宋体', size=Pt(11))

# ================= 问题一 =================
add_heading('问题一：自动化招聘筛选中的歧视')

add_q('a. 这一情形最恰当的描述是：')
add_body_mixed([('答案：（B）间接歧视。', True),
                ('平台虽移除了宗教信仰变量，但“是否可接受周末排班”与“申请休假的公共假期”这两项变量在训练数据中与宗教信仰相关，构成可识别代理变量：表面中立、实为受保护属性替代物。申请人因这些表面中立的做法被间接推断出受保护身份并受到不成比例的影响，这正是间接歧视（差异性影响）的定义。', False)])

add_q('b. 移除宗教信仰变量后差异仍然存在的正确原因是：')
add_body_mixed([('答案：（B）周末排班可接受度与休假申请这两项变量，通过与宗教信仰的相关性，依然携带着宗教信仰的影响。', True),
                ('这就是“通过无意识实现公平”的局限：移除受保护属性本身，并不能阻止代理变量把该属性的影响重新带入模型。', False)])

add_q('c. 最可能成为年龄的代理变量的选项：')
add_body_mixed([('答案：（B）距申请人简历中第一份工作至今的年数。', True),
                ('该变量与年龄高度相关（几乎随时间同步增长），是与年龄关系最直接、最明显的替代变量；就读大学、邮编、雇主数量与年龄的关联均弱于此。', False)])

add_q('d. 说明为何(c)中所选变量虽为代理变量，仍可能携带合理的预测信息：')
add_body('该变量同时反映了工作年限所代表的经验积累、技能熟练度与职业稳定性，这些因素与工作表现之间存在合法、独立的关联。即：工作年数对绩效的预测力，并非（至少并非完全）来源于它是年龄的替身，而是来源于它本身度量了与岗位相关的资历。这与邮编等“纯粹”代理变量的情形不同——邮编几乎只通过人口构成与收入水平间接起作用。因此，一个变量可以是代理变量，同时又是合理的能力信号；问题不在于它有无预测力，而在于其预测力是否独立于受保护属性。')

add_q('e. 说明为何仅移除宗教信仰（无意识公平）本身并不足够，并描述一种替代方法：')
add_body('仅移除宗教信仰不够，是因为代理性歧视机制仍然存在：周末排班可接受度、休假假期等表面中立的变量会“替”宗教信仰把影响带回模型，模型的输出依然与宗教信仰相关，歧视并未真正消除。')
add_body('本章介绍的替代方法之一是后处理的 CPV（控制受保护变量）设计，即模型 MC：先完整拟合包含宗教信仰在内的模型 M0，然后在评分时对宗教信仰取平均，Y\u0302_MC(x) = (1/N)·Σ f_M0(x_NP, X_P=x_Pj)。这样既消除了受保护属性对个体预测的直接（代理）影响，又保留了非受保护特征上的预测能力，比单纯删除变量的 MU 更彻底。')

# ================= 问题二 =================
add_heading('问题二：车险高风险标记系统（群体A与群体B各100份保单）')

# 混淆矩阵推导
add_q('推导：先由题设补全两群体混淆矩阵。')
add_body('群体A：TP=15（被标记且真实高额理赔），FN=30−15=15，FP=20−15=5，TN=100−15−5−15=65。')
add_body('群体B：TP=8，FN=10−8=2，FP=20−8=12，TN=100−8−12−2=78。')

add_q('a. 计算各群体的标记比例：')
add_body_mixed([('群体A：20/100 = ', True), ('20%', True), ('；群体B：20/100 = ', True), ('20%', True)])

add_q('b. 人口统计学均等（各群体标记比例相同）是否成立？')
add_body_mixed([('答案：', True), ('是', True), ('（20% = 20%，两群体标记比例相等，DP成立。）', False)])

add_q('c. 计算各群体的阳性预测值 PPV = Pr(Y=1 | Ŷ=1)：')
add_body_mixed([('群体A：15/20 = ', True), ('75%', True),
                ('；群体B：8/20 = ', True), ('40%', True)])

add_q('d. 充分性（各群体PPV相同）是否成立？')
add_body_mixed([('答案：', True), ('否', True), ('（75% ≠ 40%，PPV不相等，充分性被违反。）', False)])

add_q('e. 计算各群体的真阳性率 TPR = Pr(Ŷ=1 | Y=1)：')
add_body_mixed([('群体A：15/30 = ', True), ('50%', True),
                ('；群体B：8/10 = ', True), ('80%', True)])

add_q('f. 计算各群体的假阳性率 FPR = Pr(Ŷ=1 | Y=0)：')
add_body_mixed([('群体A：5/70 ≈ ', True), ('7.1%', True),
                ('；群体B：12/90 ≈ ', True), ('13.3%', True)])

add_q('g. 分离性（各群体TPR与FPR均相同）是否成立？')
add_body_mixed([('答案：', True), ('否', True),
                ('（TPR：50% vs 80%，FPR：7.1% vs 13.3%，均不相同，分离性被违反。）', False)])

add_q('h. 出现这一模式的原因是：')
add_body_mixed([('答案：（A）两个群体的真实基准率（Y=1的比例）不同。', True),
                ('群体A基准率30%，群体B基准率10%。给定模型预测能力，基准率差异会同时拉大两群体在标记比例、PPV与错误率上的差距。', False)])

add_q('i. 这对被要求“直接让模型公平”的从业者意味着什么？')
add_body_mixed([('答案：（B）应满足哪一项准则是一个事先的政策决定，仅凭更好的算法无法解决。', True),
                ('由于两群体基准率不同，分离性与充分性在数学上无法同时满足（除非完美预测或完全不利用数据），因此必须由决策者事先选定优先满足的准则。', False)])

add_q('j. 说明为何本例中人口统计学均等、分离性与充分性无法同时成立：')
add_body('记群体g的基准率 p_g = Pr(Y=1|X_P=g)。本例中 p_A=30%、p_B=10%，两者不同。由贝叶斯定理，阳性预测值 PPV_g = p_g(1−FNR_g) / [p_g(1−FNR_g)+(1−p_g)FPR_g]。')
add_body('若分离性成立（两群体 FPR、FNR 相同），则在 FPR、FNR 固定时，PPV_g 仅是 p_g 的严格增函数：对任何非完美分类器，p 越大 PPV 越大，因此 p_A≠p_B 必然导致 PPV_A≠PPV_B，充分性被违反。')
add_body('反之，若要满足充分性（PPV_A=PPV_B），就必须让两群体使用不同的 FPR/FNR，即违反分离性。本例中群体B基准率更高，若让两群体错误率相同，被标记者中B群体真实高额理赔的比例必然更高（87%式的差距）；要拉平PPV就必须对B群体施加更高门槛，从而改变其错误率。')
add_body('而人口统计学均等要求标记比例相同：本例按构造恰好同为20%，但这只是两组数据刚好相等的结果；要让标记比例相等或PPV相等，都必须为两群体设定不同的假阳性率与假阴性率，这恰恰违反分离性。因此，当基准率存在真实差异时，三项准则无法同时成立——这正是 Kleinberg 等（2017）所证明的一般性结论在本例中的体现。')

# ================= 问题三 =================
add_heading('问题三：模型设计（M0 / MU / MDP / MCDP / MC）')

add_q('a. 将下列监管干预类型与最可能满足该要求的模型设计配对：')
add_body_mixed([('i. 禁止变量规则 → ', True), ('MU（无意识模型）', True),
                ('：规则直接禁止使用受保护属性，对应FTU，属性从一开始就不作为输入。', False)])
add_body_mixed([('ii. 代理性歧视规则 → ', True), ('MC（CPV后处理）', True),
                ('：规则针对“变量未被使用但代理变量仍承载其效应”的情形；MC在评分时对受保护属性取平均，直接消除该残余影响。', False)])
add_body_mixed([('iii. 差异性影响或均等性规则 → ', True), ('MDP（人口统计学均等性）', True),
                ('：规则直接检验模型结果（如五分之四原则），要求各群体预测结果分布相同，对应DP；MDP对所有非受保护变量预处理去偏以实现DP。', False)])
add_body_mixed([('iv. 算法影响/透明度义务 → ', True), ('五种设计中的任意一种，并配合所需的流程措施', True),
                ('：这是程序性义务而非实质性义务，附加在所选设计之上，要求记录、测试与解释模型的公平性属性，不内嵌于任何一种设计。', False)])

add_q('b. MU特有的主要局限性是：')
add_body_mixed([('答案：（B）未能解决通过相关代理变量产生的间接歧视。', True),
                ('MU仅移除受保护属性本身，但代理变量仍把其影响带入模型；其余选项均不成立：MU可推广到连续结果（如回归），评分时也不需要识别受保护属性，且并非总是准确性损失最大。', False)])

add_q('c. 贷款情境下为何MCDP通常优于MDP：')
add_body('在该监管制度下，监管直接检验模型在各群体间的结果，要求结果差异仅能通过合法因素体现。收入与还款记录是合法的信用风险因素，邮编则不是。')
add_body('若用MDP，会对包括合法变量在内的全部非受保护变量去偏（正交化），这会同时抹掉收入、还款记录等合法变量带来的真实且合理的差异，导致过度纠正：一方面损害预测准确性，另一方面连监管认可的合法差异化也被一并消除，结果未必符合“仅消除非法差异”的监管意图。')
add_body('而MCDP只对非合法预测变量（邮编）去偏、保留合法变量（收入、还款记录）不变，在消除通过邮编传导的歧视性差异的同时，保留了合法的信用风险信息，恰好与“结果差异可经合法因子体现、不可经非法代理变量体现”的监管逻辑一致。因此MCDP通常优于MDP。')

# ================= 问题四 =================
add_heading('问题四：致首席精算师——车险定价模型的公平性方案备忘录')

add_body_mixed([('收件人：首席精算师　　发件人：首席数据科学家　　主题：车险定价模型上线前的公平性方案', True)])

add_body('一、建议准则：条件人口统计学均等（CDP）。本情境中受保护属性为性别，而问题二揭示两群体基准率存在真实差异。若采用无条件的DP，会强制不同风险水平的合法细分市场（如不同驾龄、理赔记录组）收取相同价格，导致逆向选择与定价扭曲；而充分性要求各群体校准一致，在定价回归情境下难以作为唯一的实质性约束。CDP在合法细分市场（监管认可的风险因子）内要求预测结果相等，既消除性别及其代理变量的影响，又允许合法风险差异存在，是公平与精算稳健性的平衡点。')
add_body('二、推荐设计：MCDP。它将预测变量分为合法变量（理赔记录、驾龄、车型等）与非合法变量（与性别相关的代理变量，如部分职业类别、居住地区等），仅对非合法变量做正交化（减去对性别回归的残差）后与合法变量一同训练。该设计在训练前去除性别经代理变量传导的影响，评分时无需再使用性别，正好落实CDP。')
add_body('三、需关注的权衡：MCDP的公平性保证依赖于“合法/非合法变量”划分的准确性与完整性；若存在未识别的不可识别代理变量，或合法变量与性别高度相关，仍可能残留间接歧视。此外，引入公平性约束会使定价与精算纯保费出现一定偏差，可能带来轻微的逆向选择风险，建议上线前用公平性—准确性权衡曲线量化该代价，并配合第四步的审计规程（预先设定测试方法与容忍度）进行验证。')

doc.save(r'C:\Users\LENOVO\Desktop\responsible-ai\TASK2\林富强_作业2.docx')
print('docx 已生成')
